import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_md_to_docx(md_path, docx_path):
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    # Split YAML front matter and body
    yaml_match = re.match(r"^---\n(.*?)\n---\n(.*)$", content, re.DOTALL)
    if yaml_match:
        yaml_text = yaml_match.group(1)
        body = yaml_match.group(2)
    else:
        body = content

    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style definitions
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b) # slate-800

    # Add Title
    title = "EpiChronos: A High-Performance Python Suite for Unified DNA Methylation and Epigenetic Aging Analysis"
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(18)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a) # slate-900

    # Add Authors
    authors_p = doc.add_paragraph()
    authors_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors_run = authors_p.add_run("Rashid Kadayil¹, Sivaranjani Chanemougame¹\n¹Department of Biotechnology, Pondicherry University, Puducherry, India")
    authors_run.font.name = 'Arial'
    authors_run.font.size = Pt(10.5)
    authors_run.italic = True
    authors_run.font.color.rgb = RGBColor(0x47, 0x55, 0x69) # slate-600

    doc.add_paragraph("\n")

    # Parse body
    lines = body.split("\n")
    in_code_block = False
    code_text = []

    for line in lines:
        stripped = line.strip()
        
        # Fenced code blocks
        if stripped.startswith("```"):
            if in_code_block:
                # End of code block
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                run = p.add_run("\n".join(code_text))
                run.font.name = 'Courier New'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                in_code_block = False
                code_text = []
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_text.append(line)
            continue

        # Headings
        if stripped.startswith("# "):
            title_text = stripped[2:].strip()
            h = doc.add_heading(level=1)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
            run = h.add_run(title_text)
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
            continue
        elif stripped.startswith("## "):
            title_text = stripped[3:].strip()
            h = doc.add_heading(level=2)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(4)
            run = h.add_run(title_text)
            run.font.name = 'Arial'
            run.font.size = Pt(12.5)
            run.bold = True
            run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
            continue
        elif stripped.startswith("### "):
            title_text = stripped[4:].strip()
            h = doc.add_heading(level=3)
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(4)
            run = h.add_run(title_text)
            run.font.name = 'Arial'
            run.font.size = Pt(11.5)
            run.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            continue

        # Empty lines
        if not stripped:
            continue

        # Bullets
        if stripped.startswith("- ") or stripped.startswith("* "):
            bullet_text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            _add_formatted_text(p, bullet_text)
            continue

        # Images/Figures
        img_match = re.match(r"^!\[(.*?)\]\((.*?)\)$", stripped)
        if img_match:
            caption = img_match.group(1)
            img_path = img_match.group(2)
            
            # Clean up pandoc label
            caption_clean = re.sub(r"\\label\{.*?\}", "", caption).strip()
            
            # Insert image
            abs_img_path = os.path.join(os.path.dirname(md_path), img_path.replace("/", os.sep))
            if os.path.exists(abs_img_path):
                try:
                    # Center-aligned paragraph for image
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.add_run().add_picture(abs_img_path, width=Inches(6.0))
                    
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.paragraph_format.space_before = Pt(4)
                    p_cap.paragraph_format.space_after = Pt(12)
                    run = p_cap.add_run(f"Figure: {caption_clean}")
                    run.font.name = 'Arial'
                    run.font.size = Pt(9.5)
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
                except Exception as e:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run(f"[Image file: {caption_clean} ({img_path})]").italic = True
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(f"[Image file: {caption_clean} ({img_path})]").italic = True
            continue

        # Math blocks
        if stripped.startswith("$$") and stripped.endswith("$$"):
            math_text = stripped[2:-2].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(math_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.italic = True
            run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
            continue

        # Standard Paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        _add_formatted_text(p, line)

def _add_formatted_text(paragraph, text):
    # Regex to extract bold **text** and inline `code` and $math$
    tokens = re.split(r"(\*\*.*?\*\*|`.*?`|\$.*?\$)", text)
    for token in tokens:
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
        elif token.startswith("$") and token.endswith("$"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(token)

if __name__ == "__main__":
    BASE = r"C:\Users\rashi\Desktop\PYTHON CODES\new 23"
    md_file = os.path.join(BASE, "paper.md")
    docx_file = os.path.join(BASE, "paper.docx")
    convert_md_to_docx(md_file, docx_file)
    print(f"Successfully generated: {docx_file}")
